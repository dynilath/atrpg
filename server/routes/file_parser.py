"""file_parser.py --- PDF / DOCX 文本提取，输出临时 TXT 到 .atrpg/uploads/"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

MAX_PARSE_CHARS = 150_000  # 单文件解析上限字符数


def parse_to_txt(file_path: Path) -> Path | None:
    """解析 PDF/DOCX → .atrpg/uploads/<同名>.txt。返回 txt 路径，失败返回 None。

    已存在同名 txt 时跳过（幂等）。
    """
    ext = file_path.suffix.lower()
    txt_path = file_path.with_suffix(".txt")
    if txt_path.exists():
        logger.debug(f"解析跳过（txt 已存在）: {txt_path.name}")
        return txt_path

    try:
        if ext == ".pdf":
            text = _parse_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = _parse_docx(file_path)
        else:
            logger.debug(f"不支持解析的格式: {ext}")
            return None

        if not text or not text.strip():
            logger.debug(f"解析结果为空: {file_path.name}")
            return None

        # 截断
        if len(text) > MAX_PARSE_CHARS:
            text = text[:MAX_PARSE_CHARS] + "\n\n[... 文本过长，已截断 ...]"

        txt_path.write_text(text, encoding="utf-8")
        logger.info(f"解析完成: {file_path.name} -> {txt_path.name} ({len(text)} chars)")
        return txt_path

    except Exception as e:
        logger.warning(f"解析失败: {file_path.name} — {e}")
        return None


def _parse_pdf(file_path: Path) -> str:
    """PyPDF2 提取 PDF 文本。"""
    from PyPDF2 import PdfReader

    reader = PdfReader(str(file_path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            content = page.extract_text() or ""
            if content.strip():
                pages.append(f"--- Page {i + 1} ---\n{content}")
        except Exception as e:
            logger.debug(f"PDF 第 {i + 1} 页读取失败: {e}")
        if sum(len(p) for p in pages) > MAX_PARSE_CHARS:
            break
    return "\n\n".join(pages)


def _parse_docx(file_path: Path) -> str:
    """python-docx 提取 DOCX 文本。"""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def list_txt_files(upload_dir: Path) -> list[dict]:
    """列出 uploads/ 下所有 txt 文件及其元数据。"""
    txts = []
    for p in sorted(upload_dir.glob("*.txt"), reverse=True):
        stat = p.stat()
        # 查找对应的原始文件
        original = None
        for ext in (".pdf", ".docx", ".doc"):
            candidate = p.with_suffix(ext)
            if candidate.exists():
                original = candidate.name
                break
        txts.append({
            "txt_name": p.name,
            "original": original,
            "size": stat.st_size,
            "chars": len(p.read_text(encoding="utf-8")),
        })
    return txts
